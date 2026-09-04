import logging
import re
from uuid import uuid4

from fastapi import APIRouter, File, HTTPException, Request, UploadFile, status
from pydantic import ValidationError

logger = logging.getLogger(__name__)

from app.core.config import settings
from app.models.content import ContentDNA, RawContent
from app.models.content_update import ContentDNAUpdate
from app.models.transformation import (
    ConflictResolution,
    DNAVersion,
    OutputGenerateRequest,
    SearchResponse,
    SearchResult,
    SourceIntegrity,
    Structure,
    StructureCreateRequest,
    StructureReferenceRequest,
    StructureSection,
    TextSourceAddRequest,
    Transformation,
    TransformationCreateRequest,
    TransformationListResponse,
    TransformationRenameRequest,
    TemplateGenerateRequest,
    URLSourceAddRequest,
    UnsupportedSourceRequest,
    WorkflowTemplate,
)

from app.services.content_dna import ContentDNAService
from app.services.ingestion import (
    IngestionError,
    PDFIngestionProvider,
    TXTIngestionProvider,
    TextIngestionProvider,
    URLIngestionProvider,
    UnsupportedIngestionProvider,
    DOCXIngestionProvider,
    YouTubeIngestionProvider,
    ImageIngestionProvider,
    AudioIngestionProvider,
    VideoIngestionProvider,
)

from app.services.llm import LLMProviderError
from app.services.output_generation import OutputGenerationService
from app.services.structure_extraction import (
    StructureExtractionError,
    StructureExtractionService,
)
from app.services.storage import (
    LocalTransformationStorage,
    TransformationNotFoundError,
)
from app.services.workflows import (
    get_workflow_templates,
    save_custom_workflow,
)
from app.services.source_integrity import (
    SourceIntegrityError,
    SourceIntegrityService,
)


router = APIRouter(
    prefix="/api/v1/transformations",
    tags=["transformations"],
)


# ============================================================
# DEPENDENCIES / HELPERS
# ============================================================


def _storage(request: Request) -> LocalTransformationStorage:
    return request.app.state.transformation_storage


def _dna_service(request: Request) -> ContentDNAService:
    return request.app.state.content_dna_service


def _outputs(request: Request) -> OutputGenerationService:
    return request.app.state.output_generation_service


def _structure_extractor(
    request: Request,
) -> StructureExtractionService:
    return request.app.state.structure_extraction_service


def _get_transformation(
    transformation_id: str,
    request: Request,
) -> Transformation:
    try:
        return _storage(request).get(transformation_id)
    except TransformationNotFoundError as exc:
        raise HTTPException(
            status_code=404,
            detail="Transformation not found",
        ) from exc


def _apply_first_source_title(
    transformation: Transformation,
    source: RawContent,
) -> Transformation:
    if (
        transformation.title == "Untitled Transformation"
        and source.title
    ):
        return transformation.model_copy(
            update={"title": source.title}
        )

    return transformation


def _merge_updates(existing: dict, updates: dict) -> dict:
    merged = existing.copy()

    for field, value in updates.items():
        if (
            isinstance(value, dict)
            and isinstance(merged.get(field), dict)
        ):
            merged[field] = _merge_updates(
                merged[field],
                value,
            )
        else:
            merged[field] = value

    return merged


def _recompute_dna(
    transformation: Transformation,
    request: Request,
) -> ContentDNA | None:
    supported_sources = [
        source
        for source in transformation.sources
        if source.metadata.get("status") != "unsupported"
        and source.text.strip()
    ]

    if not supported_sources:
        return None

    try:
        if len(supported_sources) == 1:
            return _dna_service(request).extract(
                supported_sources[0]
            )

        return _dna_service(request).extract_from_sources(
            supported_sources
        )

    except LLMProviderError as exc:
        raise HTTPException(
            status_code=502,
            detail=str(exc),
        ) from exc


def _recompute_integrity(
    transformation: Transformation,
    request: Request,
) -> SourceIntegrity:
    supported_sources = [
        source
        for source in transformation.sources
        if source.metadata.get("status") != "unsupported"
        and source.text.strip()
    ]

    if not supported_sources:
        return SourceIntegrity()

    try:
        mode = getattr(request.app.state, "llm_provider_mode", "api")
        service = SourceIntegrityService(mode=mode)
        return service.analyze(
            supported_sources,
            content_dna=transformation.content_dna,
        )
    except Exception as exc:
        logger.warning("Automatic source integrity computation failed: %s", exc)
        return transformation.source_integrity or SourceIntegrity()


def _save_dna_version(
    transformation: Transformation,
    note: str,
) -> Transformation:
    if transformation.content_dna is None:
        return transformation

    next_version = len(transformation.versions) + 1

    return transformation.model_copy(
        update={
            "versions": [
                *transformation.versions,
                DNAVersion(
                    version=next_version,
                    content_dna=transformation.content_dna,
                    note=note,
                ),
            ],
        }
    )


# ============================================================
# APPLY SOURCE-INTEGRITY RESOLUTION TO CONTENT DNA
# ============================================================


def apply_resolution_to_dna(
    dna: ContentDNA,
    claim_key: str,
    final_value: str,
    selected_claim: Claim | None = None,
    rejected_claims: list[Claim] | None = None,
) -> ContentDNA:
    """
    Apply a resolved source-integrity claim to the Content DNA.
    Actively purges all wrong/conflicted contradictory data from facts,
    findings, entities, and overview across Content DNA.
    """
    facts = dna.facts.model_copy()
    findings = dna.findings.model_copy()
    entities = dna.entities.model_copy()
    overview = dna.overview.model_copy()

    # Collect bad/rejected values and tokens
    bad_values: set[str] = set()
    if rejected_claims:
        for rc in rejected_claims:
            if rc.value and rc.value.strip():
                bad_val = rc.value.strip().lower()
                if bad_val != str(final_value).strip().lower():
                    bad_values.add(bad_val)
                # Extract significant distinctive words from the rejected claim
                stop_words = {"with", "that", "this", "from", "were", "been", "have", "some", "about", "their", "they", "among"}
                for word in re.findall(r"\b[a-z0-9]{4,}\b", bad_val):
                    if word not in str(final_value).lower() and word not in stop_words:
                        bad_values.add(word)
                # Numeric sub-tokens
                clean_nums = re.findall(r"\b\d+(?:\.\d+)?\b", rc.value)
                for num_tok in clean_nums:
                    if num_tok not in str(final_value):
                        bad_values.add(num_tok.lower())

    subj = selected_claim.subject if selected_claim and selected_claim.subject else ""
    pred = selected_claim.predicate if selected_claim and selected_claim.predicate else ""
    unit = selected_claim.unit if selected_claim and selected_claim.unit else ""
    norm_key = claim_key.replace("_", " ").lower()

    formatted_fact = (
        f"{subj} {pred}: {final_value}{' ' + unit if unit and unit not in final_value else ''}".strip()
        if (subj or pred)
        else f"{claim_key.replace('_', ' ').capitalize()}: {final_value}{' ' + unit if unit and unit not in final_value else ''}".strip()
    )

    def is_bad_entry(text: str) -> bool:
        t_lower = text.lower()
        # If matches any rejected value
        for bv in bad_values:
            if bv and bv in t_lower:
                return True
        # If starts with normalized claim key prefix but does not contain final value
        if t_lower.startswith(f"{norm_key}:") or t_lower.startswith("overall winner:") or t_lower.startswith("headquarters:") or t_lower.startswith("revenue:"):
            if final_value.lower() not in t_lower:
                return True
        return False

    # 1. Clean facts.claims
    new_claims = [item for item in facts.claims if not is_bad_entry(item)]
    if not any(final_value.lower() in c.lower() for c in new_claims):
        new_claims.append(formatted_fact)
    facts.claims = new_claims

    # 2. Clean facts.statistics
    new_stats = [item for item in facts.statistics if not is_bad_entry(item)]
    if any(ch.isdigit() for ch in final_value) and not any(final_value.lower() in s.lower() for s in new_stats):
        new_stats.append(formatted_fact)
    facts.statistics = new_stats

    # 3. Clean facts.dates
    new_dates = [item for item in facts.dates if not is_bad_entry(item)]
    if re.search(r"\b(?:19|20)\d{2}\b", final_value) and not any(final_value.lower() in d.lower() for d in new_dates):
        new_dates.append(formatted_fact)
    facts.dates = new_dates

    # 4. Clean facts.events
    facts.events = [item for item in facts.events if not is_bad_entry(item)]

    # 5. Clean findings.key_findings
    new_findings = [item for item in findings.key_findings if not is_bad_entry(item)]
    if not any(final_value.lower() in f.lower() for f in new_findings):
        new_findings.append(formatted_fact)
    findings.key_findings = new_findings

    # 6. Clean findings.risks, opportunities, implications
    findings.risks = [item for item in findings.risks if not is_bad_entry(item)]
    findings.opportunities = [item for item in findings.opportunities if not is_bad_entry(item)]
    findings.implications = [item for item in findings.implications if not is_bad_entry(item)]

    # 7. Clean entities
    if bad_values:
        entities.locations = [loc for loc in entities.locations if loc.lower() not in bad_values]
        entities.organizations = [org for org in entities.organizations if org.lower() not in bad_values]
        entities.technologies = [tech for tech in entities.technologies if tech.lower() not in bad_values]
        entities.people = [p for p in entities.people if p.lower() not in bad_values]

    if norm_key in ("headquarters", "incident location", "location", "based in", "city"):
        if final_value not in entities.locations:
            entities.locations.append(final_value)
    elif norm_key in ("cloud llm provider", "local llm", "technology"):
        if final_value not in entities.technologies:
            entities.technologies.append(final_value)

    # 8. Clean overview summary & purpose
    for bv in bad_values:
        if bv and bv in overview.summary.lower():
            pattern = re.compile(re.escape(bv), re.IGNORECASE)
            overview.summary = pattern.sub(final_value, overview.summary)
        if bv and bv in overview.purpose.lower():
            pattern = re.compile(re.escape(bv), re.IGNORECASE)
            overview.purpose = pattern.sub(final_value, overview.purpose)

    return dna.model_copy(
        update={
            "facts": facts,
            "findings": findings,
            "entities": entities,
            "overview": overview,
        }
    )


# ============================================================
# TRANSFORMATION COLLECTION
# ============================================================


@router.get(
    "",
    response_model=TransformationListResponse,
)
def list_transformations(
    request: Request,
) -> TransformationListResponse:
    return TransformationListResponse(
        transformations=_storage(request).list()
    )


@router.post(
    "",
    response_model=Transformation,
    status_code=status.HTTP_201_CREATED,
)
def create_transformation(
    payload: TransformationCreateRequest,
    request: Request,
) -> Transformation:
    transformation = Transformation(
        id=str(uuid4()),
        title=payload.title.strip()
        or "Untitled Transformation",
    )

    return _storage(request).save(transformation)


# ============================================================
# WORKFLOWS
# IMPORTANT: STATIC /workflows ROUTES MUST COME BEFORE
# /{transformation_id} ROUTES.
# ============================================================


@router.get(
    "/workflows",
    response_model=list[WorkflowTemplate],
)
def list_workflows() -> list[WorkflowTemplate]:
    return get_workflow_templates()


@router.post(
    "/workflows",
    response_model=WorkflowTemplate,
    status_code=status.HTTP_201_CREATED,
)
def create_workflow(
    payload: WorkflowTemplate,
) -> WorkflowTemplate:
    return save_custom_workflow(payload)


# ============================================================
# SINGLE TRANSFORMATION
# ============================================================


@router.get(
    "/{transformation_id}",
    response_model=Transformation,
)
def get_transformation(
    transformation_id: str,
    request: Request,
) -> Transformation:
    return _get_transformation(
        transformation_id,
        request,
    )


@router.delete(
    "/{transformation_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
def delete_transformation(
    transformation_id: str,
    request: Request,
) -> None:
    try:
        _storage(request).delete(transformation_id)
    except TransformationNotFoundError as exc:
        raise HTTPException(
            status_code=404,
            detail="Transformation not found",
        ) from exc


@router.patch(
    "/{transformation_id}/title",
    response_model=Transformation,
)
def rename_transformation(
    transformation_id: str,
    payload: TransformationRenameRequest,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    return _storage(request).save(
        transformation.model_copy(
            update={
                "title": payload.title.strip(),
            }
        )
    )


# ============================================================
# TEXT SOURCE
# ============================================================


@router.post(
    "/{transformation_id}/sources/text",
    response_model=Transformation,
    status_code=status.HTTP_201_CREATED,
)
def add_text_source(
    transformation_id: str,
    payload: TextSourceAddRequest,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    try:
        source = TextIngestionProvider().ingest(
            str(uuid4()),
            payload.title,
            payload.text,
        )
    except IngestionError as exc:
        raise HTTPException(
            status_code=422,
            detail=str(exc),
        ) from exc

    updated = transformation.model_copy(
        update={
            "sources": [
                *transformation.sources,
                source,
            ],
            "status": "processing",
        }
    )

    updated = _apply_first_source_title(
        updated,
        source,
    )

    content_dna = _recompute_dna(
        updated,
        request,
    )
    source_integrity = _recompute_integrity(
        updated,
        request,
    )

    updated = updated.model_copy(
        update={
            "content_dna": content_dna,
            "source_integrity": source_integrity,
            "status": "ready"
            if content_dna
            else "empty",
        }
    )

    if content_dna is not None:
        updated = _save_dna_version(
            updated,
            "Generated from text source",
        )

    return _storage(request).save(updated)


# ============================================================
# FILE SOURCE
# ============================================================


@router.post(
    "/{transformation_id}/sources/file",
    response_model=Transformation,
    status_code=status.HTTP_201_CREATED,
)
async def add_file_source(
    transformation_id: str,
    request: Request,
    file: UploadFile = File(...),
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    filename = file.filename or ""

    suffix = (
        filename.lower().rsplit(".", 1)[-1]
        if "." in filename
        else ""
    )

    if suffix not in {
        "txt",
        "pdf",
        "docx",
        "png",
        "jpg",
        "jpeg",
        "webp",
        "mp3",
        "wav",
        "m4a",
        "aac",
        "ogg",
        "flac",
        "wma",
        "mp4",
        "mov",
        "mkv",
        "webm",
        "avi",
    }:
        raise HTTPException(
            status_code=415,
            detail="Unsupported file type",
        )

    content = await file.read()

    if (
        len(content)
        > request.app.state.settings.max_upload_size_bytes
    ):
        raise HTTPException(
            status_code=413,
            detail="File is too large. Maximum supported file size is 256 MB.",
        )

    try:
        if suffix == "txt":
            provider = TXTIngestionProvider()

        elif suffix == "pdf":
            provider = PDFIngestionProvider()

        elif suffix == "docx":
            provider = DOCXIngestionProvider()

        elif suffix in {
            "png",
            "jpg",
            "jpeg",
            "webp",
        }:
            provider = ImageIngestionProvider()

        elif suffix in {
            "mp3",
            "wav",
            "m4a",
            "aac",
            "ogg",
            "flac",
            "wma",
        }:
            provider = AudioIngestionProvider()

        else:
            provider = VideoIngestionProvider()

        source = provider.ingest(
            str(uuid4()),
            filename,
            content,
        )

    except IngestionError as exc:
        raise HTTPException(
            status_code=422,
            detail=str(exc),
        ) from exc

    updated = transformation.model_copy(
        update={
            "sources": [
                *transformation.sources,
                source,
            ],
            "status": "processing",
        }
    )

    updated = _apply_first_source_title(
        updated,
        source,
    )

    content_dna = _recompute_dna(
        updated,
        request,
    )
    source_integrity = _recompute_integrity(
        updated,
        request,
    )

    updated = updated.model_copy(
        update={
            "content_dna": content_dna,
            "source_integrity": source_integrity,
            "status": "ready"
            if content_dna
            else "empty",
        }
    )

    if content_dna is not None:
        updated = _save_dna_version(
            updated,
            f"Generated after adding {filename}",
        )

    return _storage(request).save(updated)


# ============================================================
# URL / YOUTUBE SOURCE
# ============================================================


@router.post(
    "/{transformation_id}/sources/url",
    response_model=Transformation,
    status_code=status.HTTP_201_CREATED,
)
def add_url_source(
    transformation_id: str,
    payload: URLSourceAddRequest,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    try:
        url_lower = payload.url.lower()

        if (
            "youtube.com" in url_lower
            or "youtu.be" in url_lower
        ):
            source = YouTubeIngestionProvider().ingest(
                str(uuid4()),
                payload.title,
                payload.url,
            )
        else:
            source = URLIngestionProvider().ingest(
                str(uuid4()),
                payload.title,
                payload.url,
            )

    except IngestionError as exc:
        raise HTTPException(
            status_code=422,
            detail=str(exc),
        ) from exc

    updated = transformation.model_copy(
        update={
            "sources": [
                *transformation.sources,
                source,
            ],
            "status": "processing",
        }
    )

    updated = _apply_first_source_title(
        updated,
        source,
    )

    content_dna = _recompute_dna(
        updated,
        request,
    )
    source_integrity = _recompute_integrity(
        updated,
        request,
    )

    updated = updated.model_copy(
        update={
            "content_dna": content_dna,
            "source_integrity": source_integrity,
            "status": "ready"
            if content_dna
            else "empty",
        }
    )

    if content_dna is not None:
        updated = _save_dna_version(
            updated,
            "Generated after adding URL source",
        )

    return _storage(request).save(updated)


# ============================================================
# UNSUPPORTED SOURCE
# ============================================================


@router.post(
    "/{transformation_id}/sources/unsupported",
    response_model=Transformation,
    status_code=status.HTTP_201_CREATED,
)
def add_unsupported_source(
    transformation_id: str,
    payload: UnsupportedSourceRequest,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    source = UnsupportedIngestionProvider().ingest(
        str(uuid4()),
        payload.source_type,
        payload.title,
        payload.note,
    )

    updated = transformation.model_copy(
        update={
            "sources": [
                *transformation.sources,
                source,
            ],
        }
    )

    updated = _apply_first_source_title(
        updated,
        source,
    )

    return _storage(request).save(updated)


# ============================================================
# REMOVE SOURCE
# ============================================================


@router.delete(
    "/{transformation_id}/sources/{source_id}",
    response_model=Transformation,
)
def remove_source(
    transformation_id: str,
    source_id: str,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    sources = [
        source
        for source in transformation.sources
        if source.source_id != source_id
    ]

    if len(sources) == len(
        transformation.sources
    ):
        raise HTTPException(
            status_code=404,
            detail="Source not found",
        )

    updated = transformation.model_copy(
        update={
            "sources": sources,
            "status": "processing",
        }
    )

    content_dna = (
        _recompute_dna(
            updated,
            request,
        )
        if sources
        else None
    )
    source_integrity = (
        _recompute_integrity(
            updated,
            request,
        )
        if sources
        else SourceIntegrity()
    )

    updated = updated.model_copy(
        update={
            "content_dna": content_dna,
            "source_integrity": source_integrity,
            "status": "ready"
            if content_dna
            else "empty",
        }
    )

    if content_dna is not None:
        updated = _save_dna_version(
            updated,
            "Regenerated after removing source",
        )

    return _storage(request).save(updated)


# ============================================================
# CONTENT DNA PATCH
# ============================================================


@router.patch(
    "/{transformation_id}/content-dna",
    response_model=Transformation,
)
def patch_transformation_dna(
    transformation_id: str,
    update: ContentDNAUpdate,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    current = (
        transformation.content_dna
        or ContentDNA()
    )

    merged_data = _merge_updates(
        current.model_dump(),
        update.model_dump(exclude_unset=True),
    )

    try:
        merged_content_dna = ContentDNA.model_validate(
            merged_data
        )
    except ValidationError as exc:
        raise HTTPException(
            status_code=422,
            detail="Invalid Content DNA update",
        ) from exc

    updated = transformation.model_copy(
        update={
            "content_dna": merged_content_dna,
            "status": "ready",
        }
    )

    updated = _save_dna_version(
        updated,
        "Saved DNA edit",
    )

    return _storage(request).save(updated)


# ============================================================
# CONFLICT RESOLUTION
# ============================================================


@router.post(
    "/{transformation_id}/integrity/conflicts/{conflict_id}/resolve",
    response_model=Transformation,
)
def resolve_transformation_conflict(
    transformation_id: str,
    conflict_id: str,
    resolution: ConflictResolution,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    integrity = transformation.source_integrity
    if integrity is None:
        return transformation

    # 1. Resilient conflict matching by conflict_id, claim_key, selected_claim_id, or index
    conflict = next(
        (
            item
            for item in integrity.conflicts
            if item.conflict_id == conflict_id
            or item.conflict_id.lower() == conflict_id.lower()
        ),
        None,
    )

    if conflict is None:
        # Match by claim key
        conflict = next(
            (
                item
                for item in integrity.conflicts
                if item.claim_key == conflict_id
                or item.claim_key.lower() == conflict_id.lower()
            ),
            None,
        )

    if conflict is None and resolution.selected_claim_id:
        # Match by selected claim ID
        conflict = next(
            (
                item
                for item in integrity.conflicts
                if resolution.selected_claim_id in item.claim_ids
            ),
            None,
        )

    if conflict is None:
        # Match by numeric index (e.g. "conflict-001" -> index 0)
        try:
            clean_num = int(re.sub(r"[^0-9]", "", conflict_id)) - 1
            if 0 <= clean_num < len(integrity.conflicts):
                conflict = integrity.conflicts[clean_num]
        except (ValueError, Exception):
            pass

    if conflict is None:
        # If there is only 1 conflict in the list, safely associate with it
        if len(integrity.conflicts) == 1:
            conflict = integrity.conflicts[0]
        elif not integrity.conflicts or all(c.status == "resolved" for c in integrity.conflicts):
            # Already resolved or no pending conflicts; return current transformation cleanly
            return transformation
        else:
            raise HTTPException(
                status_code=404,
                detail=f"Conflict '{conflict_id}' not found among {len(integrity.conflicts)} conflicts",
            )

    claim_map = {
        claim.claim_id: claim
        for claim in integrity.claims
    }

    selected_claim = None

    if resolution.selected_claim_id:
        selected_claim = claim_map.get(resolution.selected_claim_id)

        if selected_claim is None:
            # Match by claim value or excerpt
            selected_claim = next(
                (
                    c
                    for c in integrity.claims
                    if c.claim_id == resolution.selected_claim_id
                    or str(c.value).strip().lower() == resolution.selected_claim_id.strip().lower()
                ),
                None,
            )

        if selected_claim is not None and selected_claim.claim_id not in conflict.claim_ids:
            conflict.claim_ids.append(selected_claim.claim_id)
    elif conflict.claim_ids:
        # Default to first conflict claim if none specified
        selected_claim = claim_map.get(conflict.claim_ids[0])

    # --------------------------------------------------------
    # Determine final value
    # --------------------------------------------------------

    if resolution.decision in {
        "accept_source_a",
        "accept_source_b",
    }:
        if selected_claim is None:
            raise HTTPException(
                status_code=400,
                detail="A claim must be selected",
            )

        final_value = str(
            selected_claim.value
        )

    elif resolution.decision == "custom_value":
        if not resolution.final_value:
            raise HTTPException(
                status_code=400,
                detail="Custom value is required",
            )

        final_value = str(
            resolution.final_value
        )

    elif resolution.decision == "retain_both":
        conflict_claims = [
            claim_map[claim_id]
            for claim_id in conflict.claim_ids
            if claim_id in claim_map
        ]

        final_value = "; ".join(
            str(claim.value)
            for claim in conflict_claims
        )

    else:
        # mark_unresolved
        final_value = None

    # --------------------------------------------------------
    # Update conflict status
    # --------------------------------------------------------

    updated_conflicts = []

    for item in integrity.conflicts:
        if item.conflict_id == conflict_id:
            updated_conflicts.append(
                item.model_copy(
                    update={
                        "status": (
                            "resolved"
                            if resolution.decision
                            != "mark_unresolved"
                            else "unresolved"
                        )
                    }
                )
            )
        else:
            updated_conflicts.append(item)

    # --------------------------------------------------------
    # Clean stale data on resolved claims
    # --------------------------------------------------------

    updated_claims = []
    for claim in integrity.claims:
        if claim.claim_id in conflict.claim_ids:
            if resolution.decision != "mark_unresolved":
                if resolution.selected_claim_id and claim.claim_id == resolution.selected_claim_id:
                    updated_claims.append(
                        claim.model_copy(
                            update={
                                "value": final_value or claim.value,
                                "status": "supported",
                            }
                        )
                    )
                elif resolution.decision == "custom_value" and claim.claim_id == conflict.claim_ids[0]:
                    updated_claims.append(
                        claim.model_copy(
                            update={
                                "value": final_value or claim.value,
                                "status": "supported",
                            }
                        )
                    )
                elif resolution.decision == "retain_both":
                    updated_claims.append(
                        claim.model_copy(
                            update={"status": "supported"}
                        )
                    )
                else:
                    # Superseded stale claim: mark resolved so it is no longer conflicting
                    updated_claims.append(
                        claim.model_copy(
                            update={"status": "resolved"}
                        )
                    )
            else:
                updated_claims.append(
                    claim.model_copy(
                        update={"status": "conflict"}
                    )
                )
        else:
            updated_claims.append(claim)

    # --------------------------------------------------------
    # Store resolution
    # --------------------------------------------------------

    updated_resolutions = [
        *integrity.resolutions,
        resolution,
    ]

    updated_integrity = integrity.model_copy(
        update={
            "claims": updated_claims,
            "conflicts": updated_conflicts,
            "resolutions": updated_resolutions,
        }
    )

    # --------------------------------------------------------
    # Apply resolved value to Content DNA & Purge Wrong Data
    # --------------------------------------------------------

    updated_dna = transformation.content_dna

    if (
        final_value is not None
        and updated_dna is not None
    ):
        rejected_claims = [
            claim_map[cid]
            for cid in conflict.claim_ids
            if cid in claim_map and (selected_claim is None or cid != selected_claim.claim_id)
        ]

        updated_dna = apply_resolution_to_dna(
            updated_dna,
            conflict.claim_key,
            final_value,
            selected_claim=selected_claim,
            rejected_claims=rejected_claims,
        )

    # --------------------------------------------------------
    # Save transformation
    # --------------------------------------------------------

    updated = transformation.model_copy(
        update={
            "source_integrity": updated_integrity,
            "content_dna": updated_dna,
            "status": (
                "ready"
                if updated_dna is not None
                else transformation.status
            ),
        }
    )

    updated = _save_dna_version(
        updated,
        f"Resolved source conflict: "
        f"{conflict.claim_key}",
    )

    return _storage(request).save(updated)


# ============================================================
# CUSTOM STRUCTURE
# ============================================================


@router.post(
    "/{transformation_id}/structures",
    response_model=Transformation,
    status_code=status.HTTP_201_CREATED,
)
def create_structure(
    transformation_id: str,
    payload: StructureCreateRequest,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    structure = Structure(
        id=str(uuid4()),
        name=payload.name.strip(),
        type=payload.type,
        source="custom",
        sections=[
            StructureSection(
                id=str(uuid4()),
                name=section.name.strip(),
                description=section.description,
                order=index,
            )
            for index, section in enumerate(
                payload.sections,
                start=1,
            )
        ],
    )

    return _storage(request).save(
        transformation.model_copy(
            update={
                "structures": [
                    *transformation.structures,
                    structure,
                ],
            }
        )
    )


# ============================================================
# REFERENCE STRUCTURE
# ============================================================


@router.post(
    "/{transformation_id}/structures/reference",
    response_model=Transformation,
    status_code=status.HTTP_201_CREATED,
)
def create_reference_structure(
    transformation_id: str,
    payload: StructureReferenceRequest,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    source = next(
        (
            item
            for item in transformation.sources
            if item.source_id
            == payload.reference_source_id
        ),
        None,
    )

    if source is None:
        raise HTTPException(
            status_code=404,
            detail="Reference source not found",
        )

    try:
        structure = _structure_extractor(
            request
        ).from_source(
            source,
            payload.name,
        )

    except StructureExtractionError as exc:
        raise HTTPException(
            status_code=422,
            detail=str(exc),
        ) from exc

    return _storage(request).save(
        transformation.model_copy(
            update={
                "structures": [
                    *transformation.structures,
                    structure,
                ],
            }
        )
    )


# ============================================================
# GENERATE OUTPUTS
# ============================================================


@router.post(
    "/{transformation_id}/outputs",
    response_model=Transformation,
)
def generate_outputs(
    transformation_id: str,
    payload: OutputGenerateRequest,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    if transformation.content_dna is None:
        if transformation.sources:
            transformation.content_dna = _llm(request).generate_content_dna(transformation.sources[0])
            _storage(request).save(transformation)
        else:
            raise HTTPException(
                status_code=409,
                detail=(
                    "Generate Content DNA before "
                    "creating outputs"
                ),
            )

    dna_version = (
        len(transformation.versions) or 1
    )

    output_types = payload.types if (payload.types or payload.structure_ids) else ["executive_summary"]

    try:
        artifacts = [
            _outputs(request).generate(
                transformation.id,
                transformation.content_dna,
                output_type,
                dna_version,
                generation_config=payload.generation_config,
            )
            for output_type in output_types
        ]

    except LLMProviderError as exc:
        raise HTTPException(
            status_code=502,
            detail=str(exc),
        ) from exc

    structure_map = {
        structure.id: structure
        for structure in transformation.structures
    }

    for structure_id in payload.structure_ids:
        structure = structure_map.get(
            structure_id
        )

        if structure is None:
            raise HTTPException(
                status_code=404,
                detail=(
                    f"Structure not found: "
                    f"{structure_id}"
                ),
            )

        artifacts.append(
            _outputs(request).generate_from_structure(
                transformation.id,
                transformation.content_dna,
                structure,
                dna_version,
            )
        )

    return _storage(request).save(
        transformation.model_copy(
            update={
                "outputs": [
                    *transformation.outputs,
                    *artifacts,
                ],
            }
        )
    )


def _parse_template_file_payload(file_base64: str | None, filename: str | None) -> tuple[str | None, str | None]:
    """Extracts template text or image base64 from an uploaded file payload (PDF, DOCX, TXT, or Image)."""
    import base64
    from io import BytesIO

    if not file_base64:
        return None, None

    raw_b64 = file_base64
    mime = ""
    if "," in file_base64 and file_base64.startswith("data:"):
        header, raw_b64 = file_base64.split(",", 1)
        mime = header.split(";")[0].replace("data:", "").lower()

    fn_lower = (filename or "").lower()

    # Image formats -> return image URL for vision parsing
    if (
        mime.startswith("image/")
        or fn_lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"))
    ):
        img_url = file_base64 if file_base64.startswith("data:image/") else f"data:image/jpeg;base64,{raw_b64}"
        return None, img_url

    try:
        file_bytes = base64.b64decode(raw_b64)
    except Exception:
        return None, None

    # PDF format
    if "pdf" in mime or fn_lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(file_bytes))
            pages = [page.extract_text() or "" for page in reader.pages]
            pdf_text = "\n\n".join(
                f"[Page {idx + 1}]\n{page_text.strip()}"
                for idx, page_text in enumerate(pages)
                if page_text.strip()
            )
            return pdf_text if pdf_text.strip() else None, None
        except Exception as exc:
            logger.warning("Failed to extract template text from PDF: %s", exc)
            return None, None

    # Word DOCX format
    if "word" in mime or "docx" in mime or fn_lower.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            parts = []
            for p in doc.paragraphs:
                txt = p.text.strip()
                if not txt:
                    continue
                style_name = (p.style.name or "").lower() if p.style else ""
                if "heading 1" in style_name:
                    parts.append(f"# {txt}")
                elif "heading 2" in style_name:
                    parts.append(f"## {txt}")
                elif "heading 3" in style_name:
                    parts.append(f"### {txt}")
                elif "list" in style_name or "bullet" in style_name:
                    parts.append(f"- {txt}")
                else:
                    parts.append(txt)

            for table in doc.tables:
                if not table.rows:
                    continue
                headers = [c.text.strip().replace("\n", " ") for c in table.rows[0].cells]
                table_md = [f"| {' | '.join(headers)} |"]
                table_md.append(f"| {' | '.join(['---'] * len(headers))} |")
                for row in table.rows[1:]:
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    table_md.append(f"| {' | '.join(cells)} |")
                parts.append("\n".join(table_md))

            docx_text = "\n\n".join(parts)
            return docx_text if docx_text.strip() else None, None
        except Exception as exc:
            logger.warning("Failed to extract template text from DOCX: %s", exc)
            return None, None

    # Plain text / Markdown
    try:
        decoded = file_bytes.decode("utf-8-sig")
        return decoded, None
    except Exception:
        return None, None


@router.post(
    "/{transformation_id}/generate-from-template",
    response_model=Transformation,
)
def generate_from_template(
    transformation_id: str,
    payload: TemplateGenerateRequest,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    if transformation.content_dna is None:
        if transformation.sources:
            transformation.content_dna = _dna_service(request).provider.generate_content_dna(transformation.sources[0])
            _storage(request).save(transformation)
        else:
            raise HTTPException(
                status_code=409,
                detail="Generate Content DNA before creating outputs",
            )

    dna_version = len(transformation.versions) or 1

    file_text, file_img = _parse_template_file_payload(
        payload.template_file_base64 or payload.template_image_base64,
        payload.template_file_name,
    )

    template_text = payload.template_text or file_text
    image_base64 = file_img or (payload.template_image_base64 if not file_text else None)

    try:
        blueprint = _outputs(request).llm_provider.extract_layout_blueprint(
            image_base64=image_base64,
            template_text=template_text,
        )

        artifact = _outputs(request).generate_from_template(
            transformation_id=transformation.id,
            content_dna=transformation.content_dna,
            blueprint=blueprint,
            dna_version=dna_version,
            template_name=payload.template_name,
            prompt=payload.user_prompt,
            generation_config=payload.generation_config,
        )

    except LLMProviderError as exc:
        raise HTTPException(
            status_code=502,
            detail=str(exc),
        ) from exc

    return _storage(request).save(
        transformation.model_copy(
            update={
                "outputs": [
                    *transformation.outputs,
                    artifact,
                ],
            }
        )
    )


@router.delete(
    "/{transformation_id}/outputs/{output_id}",
    response_model=Transformation,
)
def delete_output(
    transformation_id: str,
    output_id: str,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    remaining_outputs = [
        artifact
        for artifact in transformation.outputs
        if artifact.id != output_id
    ]

    if len(remaining_outputs) == len(transformation.outputs):
        raise HTTPException(
            status_code=404,
            detail="Output artifact not found",
        )

    return _storage(request).save(
        transformation.model_copy(
            update={
                "outputs": remaining_outputs,
            }
        )
    )


# ============================================================
# RESTORE DNA VERSION
# ============================================================


@router.post(
    "/{transformation_id}/versions/{version}/restore",
    response_model=Transformation,
)
def restore_version(
    transformation_id: str,
    version: int,
    request: Request,
) -> Transformation:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    match = next(
        (
            item
            for item in transformation.versions
            if item.version == version
        ),
        None,
    )

    if match is None:
        raise HTTPException(
            status_code=404,
            detail="DNA version not found",
        )

    updated = transformation.model_copy(
        update={
            "content_dna": match.content_dna,
            "status": "ready",
        }
    )

    updated = _save_dna_version(
        updated,
        f"Restored DNA v{version}",
    )

    return _storage(request).save(updated)


# ============================================================
# SEARCH
# ============================================================


@router.get(
    "/{transformation_id}/search",
    response_model=SearchResponse,
)
def search_transformation(
    transformation_id: str,
    q: str,
    request: Request,
) -> SearchResponse:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    needle = q.strip().lower()

    if not needle:
        return SearchResponse(
            query=q,
            results=[],
        )

    results: list[SearchResult] = []

    # --------------------------------------------------------
    # Sources
    # --------------------------------------------------------

    for source in transformation.sources:
        haystack = (
            f"{source.title}\n"
            f"{source.text}"
        )

        if needle in haystack.lower():
            results.append(
                SearchResult(
                    category="source",
                    label=source.title,
                    excerpt=source.text[:180],
                )
            )

    # --------------------------------------------------------
    # Content DNA
    # --------------------------------------------------------

    if transformation.content_dna is not None:
        for section, value in (
            transformation.content_dna
            .model_dump()
            .items()
        ):
            text = str(value)

            if needle in text.lower():
                results.append(
                    SearchResult(
                        category="dna",
                        label=section,
                        excerpt=text[:180],
                    )
                )

    # --------------------------------------------------------
    # Outputs
    # --------------------------------------------------------

    for artifact in transformation.outputs:
        if needle in artifact.content.lower():
            results.append(
                SearchResult(
                    category="output",
                    label=artifact.type,
                    excerpt=artifact.content[:180],
                )
            )

    return SearchResponse(
        query=q,
        results=results[:25],
    )


# ============================================================
# SOURCE INTEGRITY ANALYSIS
# ============================================================


@router.post(
    "/{transformation_id}/integrity",
    response_model=SourceIntegrity,
)
def analyze_source_integrity(
    transformation_id: str,
    request: Request,
) -> SourceIntegrity:
    transformation = _get_transformation(
        transformation_id,
        request,
    )

    if not transformation.sources:
        raise HTTPException(
            status_code=400,
            detail="Transformation has no sources to analyze",
        )

    try:
        mode = getattr(request.app.state, "llm_provider_mode", "api")
        service = SourceIntegrityService(
            mode=mode,
        )

        result = service.analyze(
            transformation.sources,
            content_dna=transformation.content_dna,
        )

        updated = transformation.model_copy(
            update={"source_integrity": result}
        )
        _storage(request).save(updated)

        return result

    except SourceIntegrityError as exc:
        raise HTTPException(
            status_code=422,
            detail=str(exc),
        ) from exc