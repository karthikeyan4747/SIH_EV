from io import BytesIO
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from docx import Document
from pypdf import PdfReader

from app.models.content import RawContent

import re

from youtube_transcript_api import YouTubeTranscriptApi

import cv2
import numpy as np

import subprocess
import tempfile

from app.core.config import settings

class IngestionError(ValueError):
    pass


class TextIngestionProvider:
    source_type = "text"

    def ingest(self, source_id: str, title: str, text: str) -> RawContent:
        normalized_text = text.strip()

        if not normalized_text:
            raise IngestionError("Source text cannot be empty")

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=title or "Untitled source",
            text=normalized_text,
        )


class TXTIngestionProvider:
    source_type = "txt"

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise IngestionError(
                "TXT files must be valid UTF-8 text"
            ) from exc

        return TextIngestionProvider().ingest(
            source_id,
            Path(filename).stem,
            text,
        ).model_copy(
            update={
                "source_type": self.source_type,
                "metadata": {
                    "filename": filename,
                },
            }
        )


class PDFIngestionProvider:
    source_type = "pdf"

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        try:
            reader = PdfReader(BytesIO(content))
            pages = [
                page.extract_text() or ""
                for page in reader.pages
            ]
        except Exception as exc:
            raise IngestionError(
                "Unable to read the PDF file"
            ) from exc

        text = "\n\n".join(
            f"[Page {page_number}]\n{page_text.strip()}"
            for page_number, page_text in enumerate(
                pages,
                start=1,
            )
            if page_text.strip()
        )

        if not text.strip():
            raise IngestionError(
                "The PDF does not contain extractable text"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=Path(filename).stem or "Untitled PDF",
            text=text,
            metadata={
                "filename": filename,
                "page_count": len(pages),
            },
        )


class DOCXIngestionProvider:
    source_type = "docx"

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise IngestionError(
                "Unable to read the DOCX file"
            ) from exc

        parts: list[str] = []

        # Extract normal paragraphs.
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                parts.append(text)

        # Extract tables while preserving their row/cell structure.
        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):
            rows: list[str] = []

            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                ]

                if any(cells):
                    rows.append(" | ".join(cells))

            if rows:
                parts.append(
                    f"[Table {table_index}]\n"
                    + "\n".join(rows)
                )

        text = "\n\n".join(parts).strip()

        if not text:
            raise IngestionError(
                "The DOCX file does not contain extractable text"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=Path(filename).stem or "Untitled DOCX",
            text=text,
            metadata={
                "filename": filename,
                "paragraph_count": len(
                    [
                        paragraph
                        for paragraph in document.paragraphs
                        if paragraph.text.strip()
                    ]
                ),
                "table_count": len(document.tables),
            },
        )


class ImageIngestionProvider:
    source_type = "image"

    def __init__(self) -> None:
        from rapidocr_onnxruntime import RapidOCR

        self.ocr = RapidOCR()


    def extract_text_from_image_array(self, image_array) -> str:
        try:
            result, _ = self.ocr(image_array)
        except Exception as exc:
            raise IngestionError(
                "Unable to analyze a video frame"
            ) from exc

        if not result:
            return ""

        parts: list[str] = []

        for item in result:
            text = str(item[1]).strip()

            if text:
                parts.append(text)

        return "\n".join(parts).strip()

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        try:
            image_array = cv2.imdecode(
                np.frombuffer(content, dtype=np.uint8),
                cv2.IMREAD_COLOR,
            )

            if image_array is None:
                raise IngestionError(
                    "Unable to decode the image file"
                )

            result, _ = self.ocr(image_array)

        except IngestionError:
            raise
        except Exception as exc:
            raise IngestionError(
                "Unable to process the image"
            ) from exc

        if not result:
            raise IngestionError(
                "No readable text was detected in the image"
            )

        text_parts: list[str] = []

        for item in result:
            text = str(item[1]).strip()

            if text:
                text_parts.append(text)

        text = "\n".join(text_parts).strip()

        if not text:
            raise IngestionError(
                "No readable text was detected in the image"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=Path(filename).stem or "Untitled image",
            text=text,
            metadata={
                "filename": filename,
                "ocr_engine": "RapidOCR",
            },
        )


class AudioIngestionProvider:
    source_type = "audio"

    def __init__(self) -> None:
        from faster_whisper import WhisperModel

        self.model = WhisperModel(
            "base",
            device="cpu",
            compute_type="int8",
        )

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        temp_path: Path | None = None

        try:
            suffix = Path(filename).suffix or ".audio"

            import tempfile

            with tempfile.NamedTemporaryFile(
                delete=False,
                suffix=suffix,
            ) as temp_file:
                temp_file.write(content)
                temp_path = Path(temp_file.name)

            segments, info = self.model.transcribe(
                str(temp_path),
                beam_size=5,
                vad_filter=True,
            )

            text_parts: list[str] = []

            for segment in segments:
                text = segment.text.strip()

                if text:
                    text_parts.append(text)

            text = "\n".join(text_parts).strip()

            if not text:
                raise IngestionError(
                    "No speech could be transcribed from the audio"
                )

            return RawContent(
                source_id=source_id,
                source_type=self.source_type,
                title=Path(filename).stem or "Untitled audio",
                text=text,
                metadata={
                    "filename": filename,
                    "transcription_engine": "faster-whisper",
                    "model": "base",
                    "language": getattr(info, "language", ""),
                    "language_probability": getattr(
                        info,
                        "language_probability",
                        None,
                    ),
                    "duration_seconds": getattr(
                        info,
                        "duration",
                        None,
                    ),
                },
            )

        except IngestionError:
            raise

        except Exception as exc:
            raise IngestionError(
                "Unable to transcribe the audio file"
            ) from exc

        finally:
            if temp_path is not None:
                temp_path.unlink(missing_ok=True)


def extract_video_frames(
    video_path: Path,
    interval_seconds: float = 5.0,
) -> list[tuple[float, object]]:
    capture = cv2.VideoCapture(str(video_path))

    if not capture.isOpened():
        raise IngestionError("Unable to open the video file")

    fps = capture.get(cv2.CAP_PROP_FPS)

    if not fps or fps <= 0:
        capture.release()
        raise IngestionError("Unable to determine video frame rate")

    frame_count = capture.get(cv2.CAP_PROP_FRAME_COUNT)
    duration = frame_count / fps if frame_count else 0

    frames: list[tuple[float, object]] = []
    current_time = 0.0

    while current_time <= duration:
        capture.set(
            cv2.CAP_PROP_POS_MSEC,
            current_time * 1000,
        )

        success, frame = capture.read()

        if success and frame is not None:
            frames.append((current_time, frame))

        current_time += interval_seconds

    capture.release()

    return frames

class VideoIngestionProvider:
    source_type = "video"

    def __init__(self) -> None:
        self.audio_provider = AudioIngestionProvider()
        self.image_provider = ImageIngestionProvider()

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        video_path: Path | None = None
        audio_path: Path | None = None

        try:
            video_suffix = Path(filename).suffix or ".video"

            # Save uploaded video temporarily.
            with tempfile.NamedTemporaryFile(
                delete=False,
                suffix=video_suffix,
            ) as video_file:
                video_file.write(content)
                video_path = Path(video_file.name)

            # -----------------------------
            # VISUAL BRANCH
            # -----------------------------
            frames = extract_video_frames(
                video_path,
                interval_seconds=5.0,
            )

            visual_parts: list[str] = []

            for timestamp, frame in frames:
                frame_text = (
                    self.image_provider.extract_text_from_image_array(
                        frame
                    )
                )

                if frame_text:
                    visual_parts.append(
                        f"[Visual frame at {timestamp:.1f}s]\n"
                        f"{frame_text}"
                    )

            # -----------------------------
            # AUDIO BRANCH
            # -----------------------------
            with tempfile.NamedTemporaryFile(
                delete=False,
                suffix=".wav",
            ) as audio_file:
                audio_path = Path(audio_file.name)

            result = subprocess.run(
                [
                    settings.ffmpeg_path,
                    "-y",
                    "-i",
                    str(video_path),
                    "-vn",
                    "-ac",
                    "1",
                    "-ar",
                    "16000",
                    "-c:a",
                    "pcm_s16le",
                    str(audio_path),
                ],
                capture_output=True,
                text=True,
                timeout=300,
            )

            if result.returncode != 0:
                raise IngestionError(
                    "Unable to extract audio from the video file"
                )

            if (
                not audio_path.exists()
                or audio_path.stat().st_size == 0
            ):
                raise IngestionError(
                    "The video does not contain an extractable audio track"
                )

            audio_content = audio_path.read_bytes()

            audio_result = self.audio_provider.ingest(
                source_id=source_id,
                filename=f"{Path(filename).stem}.wav",
                content=audio_content,
            )

            # -----------------------------
            # COMBINE AUDIO + VISUAL TEXT
            # -----------------------------
            combined_text = audio_result.text

            if visual_parts:
                combined_text += (
                    "\n\n"
                    + "\n\n".join(visual_parts)
                )

            return RawContent(
                source_id=source_id,
                source_type=self.source_type,
                title=(
                    Path(filename).stem
                    or "Untitled video"
                ),
                text=combined_text,
                metadata={
                    "filename": filename,
                    "transcription_engine": audio_result.metadata.get(
                        "transcription_engine",
                        "faster-whisper",
                    ),
                    "model": audio_result.metadata.get(
                        "model",
                        "base",
                    ),
                    "language": audio_result.metadata.get(
                        "language",
                        "",
                    ),
                    "language_probability": audio_result.metadata.get(
                        "language_probability",
                        None,
                    ),
                    "duration_seconds": audio_result.metadata.get(
                        "duration_seconds",
                        None,
                    ),
                    "audio_extracted": True,
                    "visual_analysis": bool(visual_parts),
                    "ocr_engine": (
                        "RapidOCR"
                        if visual_parts
                        else ""
                    ),
                    "frame_interval_seconds": 5.0,
                    "visual_frame_count": len(frames),
                    "ffmpeg": settings.ffmpeg_path,
                },
            )

        except IngestionError:
            raise

        except subprocess.TimeoutExpired as exc:
            raise IngestionError(
                "Video audio extraction timed out"
            ) from exc

        except Exception as exc:
            raise IngestionError(
                "Unable to process the video file"
            ) from exc

        finally:
            if video_path is not None:
                video_path.unlink(
                    missing_ok=True
                )

            if audio_path is not None:
                audio_path.unlink(
                    missing_ok=True
                )
class URLIngestionProvider:
    source_type = "url"

    def ingest(
        self,
        source_id: str,
        title: str,
        url: str,
    ) -> RawContent:
        parsed = urlparse(url)

        if (
            parsed.scheme not in {"http", "https"}
            or not parsed.netloc
        ):
            raise IngestionError(
                "Enter a valid HTTP or HTTPS URL"
            )

        request = Request(
            url,
            headers={
                "User-Agent": "EV Content Transformation/0.1"
            },
        )

        try:
            with urlopen(
                request,
                timeout=10,
            ) as response:
                content_type = response.headers.get(
                    "content-type",
                    "",
                )

                if (
                    "text/html" not in content_type
                    and "text/plain" not in content_type
                ):
                    raise IngestionError(
                        "This URL does not expose readable text content"
                    )

                raw = response.read(
                    1_000_000
                ).decode(
                    "utf-8",
                    errors="replace",
                )

        except HTTPError as exc:
            raise IngestionError(
                f"URL could not be fetched: HTTP {exc.code}"
            ) from exc

        except URLError as exc:
            raise IngestionError(
                "URL could not be fetched"
            ) from exc

        text = " ".join(
            raw.replace("<", " <")
            .replace(">", "> ")
            .split()
        )

        if not text.strip():
            raise IngestionError(
                "No readable text could be extracted from this URL"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=title.strip() or parsed.netloc,
            text=text[:100_000],
            metadata={
                "url": url,
            },
        )

class YouTubeIngestionProvider:
    source_type = "youtube"

    def ingest(
        self,
        source_id: str,
        title: str,
        url: str,
    ) -> RawContent:
        video_id = self._extract_video_id(url)

        if not video_id:
            raise IngestionError(
                "Enter a valid YouTube video URL"
            )

        try:
            transcript = YouTubeTranscriptApi().fetch(
                video_id,
                languages=["en", "ta", "hi"],
            )
        except Exception as exc:
            raise IngestionError(
                f"YouTube transcript retrieval failed: {type(exc).__name__}: {exc}"
            ) from exc

        snippets = getattr(transcript, "snippets", [])

        text = "\n".join(
            snippet.text.strip()
            for snippet in snippets
            if getattr(snippet, "text", "").strip()
        ).strip()

        if not text:
            raise IngestionError(
                "The YouTube video does not contain an accessible transcript"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=title.strip() or "YouTube source",
            text=text,
            metadata={
                "url": url,
                "video_id": video_id,
                "transcript_language": getattr(
                    transcript,
                    "language",
                    "",
                ),
                "transcript_language_code": getattr(
                    transcript,
                    "language_code",
                    "",
                ),
                "is_generated": getattr(
                    transcript,
                    "is_generated",
                    None,
                ),
            },
        )

    @staticmethod
    def _extract_video_id(url: str) -> str:
        patterns = [
            r"(?:v=)([A-Za-z0-9_-]{11})",
            r"(?:youtu\.be/)([A-Za-z0-9_-]{11})",
            r"(?:youtube\.com/shorts/)([A-Za-z0-9_-]{11})",
            r"(?:youtube\.com/embed/)([A-Za-z0-9_-]{11})",
            r"(?:youtube\.com/live/)([A-Za-z0-9_-]{11})",
        ]

        for pattern in patterns:
            match = re.search(pattern, url)

            if match:
                return match.group(1)

        return ""




class UnsupportedIngestionProvider:
    def ingest(
        self,
        source_id: str,
        source_type: str,
        title: str,
        note: str = "",
    ) -> RawContent:
        return RawContent(
            source_id=source_id,
            source_type=source_type,
            title=(
                title.strip()
                or f"Unsupported {source_type.upper()} source"
            ),
            text=(
                f"{source_type.upper()} processing is not available "
                "in this EV backend configuration. This source was "
                "recorded for provenance but was not used to generate "
                "Content DNA."
            ),
            metadata={
                "status": "unsupported",
                "note": note,
            },
        )