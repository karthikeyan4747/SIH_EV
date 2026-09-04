from pathlib import Path
import pytest
from fastapi.testclient import TestClient

from app.main import app
from app.models.content import ContentDNA, Identity, Overview, Facts, Findings, Recommendations, Context, Evidence, RawContent
from app.models.transformation import Transformation, Artifact
from app.services.llm import _extract_blueprint_from_text, _deterministic_generate_from_blueprint
from app.services.storage import LocalTransformationStorage


class FakeTemplateProvider:
    def generate_content_dna(self, content: RawContent) -> ContentDNA:
        return ContentDNA(
            identity=Identity(title=content.title, content_type="text", source_description="Quarterly EV report"),
            overview=Overview(summary="Comprehensive roadmap deploying 15,000 EV charging hubs across 42 cities with $450M.", purpose="Planning"),
            facts=Facts(
                claims=["Project Apollo aims to deploy 15,000 EV hubs across 42 cities.", "Capital allocation is set at $450 Million."],
                statistics=["15,000 charging hubs", "$450 Million budget", "42 metropolitan cities"],
            ),
            findings=Findings(
                key_findings=["Grid capacity constraints in Tier-2 regions remain primary deployment bottleneck."],
                risks=["Supply chain delays on DC fast-charging transformers."],
            ),
            recommendations=Recommendations(
                recommendations=["Partner with regional utility grids for substation capacity reservations.", "Implement dynamic load balancing."],
            ),
            context=Context(target_audience="Board of Directors", communication_objective="Approve Phase 2 budget"),
            evidence=Evidence(supporting_excerpt="Project Apollo aims to deploy 15,000 EV hubs across 42 cities with a $450M allocation."),
        )

    def generate_output(
        self,
        content_dna: ContentDNA,
        output_type: str,
        output_spec: dict,
        user_prompt: str | None = None,
        generation_config: dict | None = None,
    ) -> str:
        return f"# Generated {output_type}\n\nContent for {content_dna.identity.title}"

    def extract_layout_blueprint(
        self,
        image_base64: str | None = None,
        template_text: str | None = None,
    ) -> dict:
        if template_text:
            return _extract_blueprint_from_text(template_text)
        return {
            "title": "Executive Board Briefing Clone",
            "layout_type": "2_column",
            "sections": [
                {"heading": "Key Performance Metrics", "style": "kpi_cards", "description": "Top metrics"},
                {"heading": "Strategic Directives", "style": "bullet_list", "description": "Actions"},
            ],
        }

    def generate_output_from_template(
        self,
        content_dna: ContentDNA,
        blueprint: dict,
        user_prompt: str | None = None,
        generation_config: dict | None = None,
    ) -> str:
        return _deterministic_generate_from_blueprint(content_dna, blueprint, generation_config)


@pytest.fixture
def client(tmp_path: Path) -> TestClient:
    app.state.transformation_storage = LocalTransformationStorage(str(tmp_path / "transformations.json"))
    fake = FakeTemplateProvider()
    app.state.content_dna_service.provider = fake
    app.state.output_generation_service.llm_provider = fake
    return TestClient(app)


def test_extract_blueprint_from_text():
    sample_template = """# Q3 Board Strategic Memo
## Executive Synthesis
Overview of progress and goals.

## Key Performance Metrics
- Metric 1
- Metric 2

## Risk & Mitigation Matrix
| Risk | Mitigation | Impact |

## Strategic Action Checklist
1. Step one
2. Step two
"""
    blueprint = _extract_blueprint_from_text(sample_template)
    assert blueprint["title"] == "Q3 Board Strategic Memo"
    assert len(blueprint["sections"]) >= 4
    styles = [s["style"] for s in blueprint["sections"]]
    assert "kpi_cards" in styles or "bullet_list" in styles
    assert "table" in styles


def test_deterministic_generate_from_blueprint():
    dna = FakeTemplateProvider().generate_content_dna(RawContent(source_id="s1", title="EV", text="EV", source_type="text"))
    blueprint = {
        "title": "McKinsey Style 1-Page Executive Memo",
        "sections": [
            {"heading": "Executive Summary", "style": "paragraph", "description": "High level synopsis"},
            {"heading": "Key Performance Metrics", "style": "kpi_cards", "description": "KPI callout cards"},
            {"heading": "Strategic Directives", "style": "bullet_list", "description": "Actionable items"},
            {"heading": "Risk & Discrepancy Matrix", "style": "table", "description": "Audit matrix"},
        ],
    }
    output = _deterministic_generate_from_blueprint(dna, blueprint, {"audience": "C-Suite", "tone": "Decisive"})
    assert "# McKinsey Style 1-Page Executive Memo" in output
    assert "15,000 charging hubs" in output or "450 Million" in output
    assert "## 1. Executive Summary" in output
    assert "## 2. Key Performance Metrics" in output
    assert "## 3. Strategic Directives" in output
    assert "## 4. Risk & Discrepancy Matrix" in output


def test_generate_from_template_endpoint(client: TestClient):
    # 1. Create a transformation
    resp = client.post("/api/v1/transformations", json={"title": "Template Clone Test"})
    assert resp.status_code == 201
    trans_id = resp.json()["id"]

    # 2. Ingest a text source (automatically computes content_dna)
    src_resp = client.post(
        f"/api/v1/transformations/{trans_id}/sources/text",
        json={"title": "EV Strategy 2027", "text": "Project Apollo aims to deploy 15,000 EV hubs across 42 cities with a $450 Million budget. Grid capacity in Tier-2 regions is the main bottleneck."},
    )
    assert src_resp.status_code in (200, 201)
    trans_data = src_resp.json()
    assert trans_data["content_dna"] is not None

    # 3. Generate Cloned Deliverable from Template Text & Image
    fake_png_base64 = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="

    template_payload = {
        "template_name": "Executive Board Briefing Clone",
        "template_image_base64": fake_png_base64,
        "template_text": "# Executive Board Memo\n## Key Performance Metrics\n## Strategic Findings\n## Action Plan",
        "generation_config": {"audience": "Executive Leadership", "tone": "Decisive"},
    }

    gen_resp = client.post(
        f"/api/v1/transformations/{trans_id}/generate-from-template",
        json=template_payload,
    )
    assert gen_resp.status_code == 200
    data = gen_resp.json()
    assert len(data["outputs"]) >= 1
    cloned_output = [o for o in data["outputs"] if o["type"] == "template_clone"]
    assert len(cloned_output) == 1
    assert cloned_output[0]["structure_id"] == "Executive Board Briefing Clone"
    assert len(cloned_output[0]["content"]) > 50


def test_generate_from_docx_template_file(client: TestClient):
    import base64
    from io import BytesIO
    from docx import Document

    # Create transformation
    resp = client.post("/api/v1/transformations", json={"title": "DOCX Template Test"})
    assert resp.status_code == 201
    trans_id = resp.json()["id"]

    # Ingest text source
    src_resp = client.post(
        f"/api/v1/transformations/{trans_id}/sources/text",
        json={"title": "Data", "text": "Apollo deployed 15,000 hubs with $450 Million."},
    )
    assert src_resp.status_code in (200, 201)

    # Create a DOCX in memory
    doc = Document()
    doc.add_heading("Board Strategic Report Format", level=1)
    doc.add_paragraph("High level overview of operational performance.")
    doc.add_heading("Key Metrics Matrix", level=2)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Objective"
    table.rows[0].cells[1].text = "Status"
    table.rows[1].cells[0].text = "Hub Deployment"
    table.rows[1].cells[1].text = "On Track"
    docx_io = BytesIO()
    doc.save(docx_io)
    docx_b64 = base64.b64encode(docx_io.getvalue()).decode("utf-8")

    gen_resp = client.post(
        f"/api/v1/transformations/{trans_id}/generate-from-template",
        json={
            "template_name": "DOCX Report Clone",
            "template_file_name": "reference_format.docx",
            "template_file_base64": f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{docx_b64}",
            "generation_config": {"tone": "Executive"},
        },
    )
    assert gen_resp.status_code == 200
    data = gen_resp.json()
    outputs = [o for o in data["outputs"] if o["type"] == "template_clone"]
    assert len(outputs) == 1
    assert "Board Strategic Report Format" in outputs[0]["content"] or "Key Metrics Matrix" in outputs[0]["content"] or "Executive" in outputs[0]["content"]
