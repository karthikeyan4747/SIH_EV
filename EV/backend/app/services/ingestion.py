from io import BytesIO
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen

from docx import Document
from pypdf import PdfReader

from app.models.content import RawContent

import re

from youtube_transcript_api import YouTubeTranscriptApi

from rapidocr_onnxruntime import RapidOCR
import cv2
import numpy as np


from faster_whisper import WhisperModel

class IngestionError(ValueError):
    pass


class TextIngestionProvider:
    source_type = "text"

    def ingest(self, source_id: str, title: str, text: str) -> RawContent:
        normalized_text = text.strip()

        if not normalized_text:
            raise IngestionError("Source text cannot be empty")

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=title.strip() or "Untitled source",
            text=normalized_text,
        )


class TXTIngestionProvider:
    source_type = "txt"

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise IngestionError(
                "TXT files must be valid UTF-8 text"
            ) from exc

        return TextIngestionProvider().ingest(
            source_id,
            Path(filename).stem,
            text,
        ).model_copy(
            update={
                "source_type": self.source_type,
                "metadata": {
                    "filename": filename,
                },
            }
        )


class PDFIngestionProvider:
    source_type = "pdf"

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        try:
            reader = PdfReader(BytesIO(content))
            pages = [
                page.extract_text() or ""
                for page in reader.pages
            ]
        except Exception as exc:
            raise IngestionError(
                "Unable to read the PDF file"
            ) from exc

        text = "\n\n".join(
            f"[Page {page_number}]\n{page_text.strip()}"
            for page_number, page_text in enumerate(
                pages,
                start=1,
            )
            if page_text.strip()
        )

        if not text.strip():
            raise IngestionError(
                "The PDF does not contain extractable text"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=Path(filename).stem or "Untitled PDF",
            text=text,
            metadata={
                "filename": filename,
                "page_count": len(pages),
            },
        )


class DOCXIngestionProvider:
    source_type = "docx"

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise IngestionError(
                "Unable to read the DOCX file"
            ) from exc

        parts: list[str] = []

        # Extract normal paragraphs.
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                parts.append(text)

        # Extract tables while preserving their row/cell structure.
        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):
            rows: list[str] = []

            for row in table.rows:
                cells = [
                    cell.text.strip().replace("\n", " ")
                    for cell in row.cells
                ]

                if any(cells):
                    rows.append(" | ".join(cells))

            if rows:
                parts.append(
                    f"[Table {table_index}]\n"
                    + "\n".join(rows)
                )

        text = "\n\n".join(parts).strip()

        if not text:
            raise IngestionError(
                "The DOCX file does not contain extractable text"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=Path(filename).stem or "Untitled DOCX",
            text=text,
            metadata={
                "filename": filename,
                "paragraph_count": len(
                    [
                        paragraph
                        for paragraph in document.paragraphs
                        if paragraph.text.strip()
                    ]
                ),
                "table_count": len(document.tables),
            },
        )


class ImageIngestionProvider:
    source_type = "image"

    def __init__(self) -> None:
        self.ocr = RapidOCR()

    def ingest(
        self,
        source_id: str,
        filename: str,
        content: bytes,
    ) -> RawContent:
        try:
            image_array = cv2.imdecode(
                np.frombuffer(content, dtype=np.uint8),
                cv2.IMREAD_COLOR,
            )

            if image_array is None:
                raise IngestionError(
                    "Unable to decode the image file"
                )

            result, _ = self.ocr(image_array)

        except IngestionError:
            raise
        except Exception as exc:
            raise IngestionError(
                "Unable to process the image"
            ) from exc

        if not result:
            raise IngestionError(
                "No readable text was detected in the image"
            )

        text_parts: list[str] = []

        for item in result:
            text = str(item[1]).strip()

            if text:
                text_parts.append(text)

        text = "\n".join(text_parts).strip()

        if not text:
            raise IngestionError(
                "No readable text was detected in the image"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=Path(filename).stem or "Untitled image",
            text=text,
            metadata={
                "filename": filename,
                "ocr_engine": "RapidOCR",
            },
        )


class URLIngestionProvider:
    source_type = "url"

    def ingest(
        self,
        source_id: str,
        title: str,
        url: str,
    ) -> RawContent:
        parsed = urlparse(url)

        if (
            parsed.scheme not in {"http", "https"}
            or not parsed.netloc
        ):
            raise IngestionError(
                "Enter a valid HTTP or HTTPS URL"
            )

        request = Request(
            url,
            headers={
                "User-Agent": "EV Content Transformation/0.1"
            },
        )

        try:
            with urlopen(
                request,
                timeout=10,
            ) as response:
                content_type = response.headers.get(
                    "content-type",
                    "",
                )

                if (
                    "text/html" not in content_type
                    and "text/plain" not in content_type
                ):
                    raise IngestionError(
                        "This URL does not expose readable text content"
                    )

                raw = response.read(
                    1_000_000
                ).decode(
                    "utf-8",
                    errors="replace",
                )

        except HTTPError as exc:
            raise IngestionError(
                f"URL could not be fetched: HTTP {exc.code}"
            ) from exc

        except URLError as exc:
            raise IngestionError(
                "URL could not be fetched"
            ) from exc

        text = " ".join(
            raw.replace("<", " <")
            .replace(">", "> ")
            .split()
        )

        if not text.strip():
            raise IngestionError(
                "No readable text could be extracted from this URL"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=title.strip() or parsed.netloc,
            text=text[:100_000],
            metadata={
                "url": url,
            },
        )

class YouTubeIngestionProvider:
    source_type = "youtube"

    def ingest(
        self,
        source_id: str,
        title: str,
        url: str,
    ) -> RawContent:
        video_id = self._extract_video_id(url)

        if not video_id:
            raise IngestionError(
                "Enter a valid YouTube video URL"
            )

        try:
            transcript = YouTubeTranscriptApi().fetch(
                video_id,
                languages=["en", "ta", "hi"],
            )
        except Exception as exc:
            raise IngestionError(
                f"YouTube transcript retrieval failed: {type(exc).__name__}: {exc}"
            ) from exc

        snippets = getattr(transcript, "snippets", [])

        text = "\n".join(
            snippet.text.strip()
            for snippet in snippets
            if getattr(snippet, "text", "").strip()
        ).strip()

        if not text:
            raise IngestionError(
                "The YouTube video does not contain an accessible transcript"
            )

        return RawContent(
            source_id=source_id,
            source_type=self.source_type,
            title=title.strip() or "YouTube source",
            text=text,
            metadata={
                "url": url,
                "video_id": video_id,
                "transcript_language": getattr(
                    transcript,
                    "language",
                    "",
                ),
                "transcript_language_code": getattr(
                    transcript,
                    "language_code",
                    "",
                ),
                "is_generated": getattr(
                    transcript,
                    "is_generated",
                    None,
                ),
            },
        )

    @staticmethod
    def _extract_video_id(url: str) -> str:
        patterns = [
            r"(?:v=)([A-Za-z0-9_-]{11})",
            r"(?:youtu\.be/)([A-Za-z0-9_-]{11})",
            r"(?:youtube\.com/shorts/)([A-Za-z0-9_-]{11})",
            r"(?:youtube\.com/embed/)([A-Za-z0-9_-]{11})",
            r"(?:youtube\.com/live/)([A-Za-z0-9_-]{11})",
        ]

        for pattern in patterns:
            match = re.search(pattern, url)

            if match:
                return match.group(1)

        return ""




class UnsupportedIngestionProvider:
    def ingest(
        self,
        source_id: str,
        source_type: str,
        title: str,
        note: str = "",
    ) -> RawContent:
        return RawContent(
            source_id=source_id,
            source_type=source_type,
            title=(
                title.strip()
                or f"Unsupported {source_type.upper()} source"
            ),
            text=(
                f"{source_type.upper()} processing is not available "
                "in this EV backend configuration. This source was "
                "recorded for provenance but was not used to generate "
                "Content DNA."
            ),
            metadata={
                "status": "unsupported",
                "note": note,
            },
        )